import io
import struct
import unittest
import zlib

from docx import Document
from docx.shared import Inches

from app.routers.document import _is_feature_sized, _is_logo_like, _select_logo
from app.services.doc_parser import DocImage, OutlineBlock, ParsedDocument, parse_docx
from app.services.doc_structure import DocImageRef, split_into_pages


def _png(w: int, h: int, rgb=(10, 20, 30)) -> bytes:
    raw = b"".join(b"\x00" + bytes(rgb) * w for _ in range(h))

    def chunk(tag: bytes, data: bytes) -> bytes:
        body = tag + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)

    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


def _parsed_about_doc() -> ParsedDocument:
    outline = [
        OutlineBlock(1, "Acme"),
        OutlineBlock(0, "Welcome."),
        OutlineBlock(1, "About Us"),
        OutlineBlock(0, "Our history."),
    ]
    return ParsedDocument(
        source_kind="docx",
        source_ref="a.docx",
        title="Acme",
        raw_text="\n".join(b.text for b in outline),
        headings=[],
        outline=outline,
    )


class ImagePlacementTest(unittest.TestCase):
    def test_image_attaches_to_page_with_heading_alt(self):
        parsed = _parsed_about_doc()
        # Anchored right after the "About Us" heading (outline index 2 → anchor 3).
        refs = [DocImageRef(url="data:about", width=400, height=300, anchor=3)]
        sc = split_into_pages(parsed, images=refs)

        self.assertEqual(sc.images, [])  # homepage got none
        about = sc.discovered_pages[0]
        self.assertEqual([m.url for m in about.image_metadata], ["data:about"])
        self.assertEqual(about.image_metadata[0].alt, "About Us")
        self.assertEqual(about.image_metadata[0].intent, "hero")

    def test_image_before_any_heading_is_home_hero(self):
        parsed = _parsed_about_doc()
        refs = [DocImageRef(url="data:cover", width=900, height=600, anchor=0)]
        sc = split_into_pages(parsed, images=refs)
        self.assertEqual([m.url for m in sc.image_metadata], ["data:cover"])
        self.assertEqual(sc.image_metadata[0].intent, "hero")

    def test_second_image_on_a_page_is_generic(self):
        parsed = _parsed_about_doc()
        refs = [
            DocImageRef(url="d1", width=400, height=300, anchor=3),
            DocImageRef(url="d2", width=400, height=300, anchor=4),
        ]
        sc = split_into_pages(parsed, images=refs)
        about = sc.discovered_pages[0]
        self.assertEqual([m.intent for m in about.image_metadata], ["hero", "generic"])


class LogoAndFilterTest(unittest.TestCase):
    def test_logo_like_detection(self):
        self.assertTrue(_is_logo_like(DocImage(b"", "image/png", width=200, height=80)))
        self.assertFalse(_is_logo_like(DocImage(b"", "image/png", width=1200, height=800)))
        self.assertFalse(_is_logo_like(DocImage(b"", "image/png")))  # unknown dims

    def test_select_logo_picks_small_cover_image(self):
        imgs = [
            DocImage(b"", "image/png", width=180, height=60, anchor=0),
            DocImage(b"", "image/png", width=800, height=600, anchor=4),
        ]
        self.assertIs(_select_logo(imgs), imgs[0])

    def test_select_logo_skips_large_hero_first(self):
        self.assertIsNone(_select_logo([DocImage(b"", "image/png", width=1200, height=800)]))

    def test_feature_size_filter_drops_icons_keeps_photos(self):
        self.assertTrue(_is_feature_sized(DocImage(b"", "image/png", width=400, height=300)))
        self.assertFalse(_is_feature_sized(DocImage(b"", "image/png", width=64, height=64)))
        self.assertTrue(_is_feature_sized(DocImage(b"", "image/png")))  # unknown dims pass


class DocxImageExtractionTest(unittest.TestCase):
    def test_inline_image_captures_dims_and_position(self):
        d = Document()
        d.add_heading("Acme", 0)
        d.add_heading("About Us", 1)
        d.add_paragraph("History.")
        d.add_picture(io.BytesIO(_png(320, 240)), width=Inches(2))
        d.add_heading("Contact", 1)
        buf = io.BytesIO()
        d.save(buf)

        parsed = parse_docx(buf.getvalue(), "a.docx")
        self.assertEqual(len(parsed.images), 1)
        img = parsed.images[0]
        self.assertEqual((img.width, img.height), (320, 240))
        # Sits after the About section's copy, before the Contact heading.
        self.assertEqual(parsed.outline[img.anchor - 1].text, "History.")

        # End-to-end: the image lands on the About page with About-derived alt.
        from app.services.doc_structure import DocImageRef as Ref

        sc = split_into_pages(
            parsed,
            images=[Ref(url="data:x", width=img.width, height=img.height, anchor=img.anchor)],
        )
        about = next(p for p in sc.discovered_pages if p.url_path == "/about-us")
        self.assertEqual(about.image_metadata[0].alt, "About Us")


if __name__ == "__main__":
    unittest.main()
