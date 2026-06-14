"""
PDF + DOCX → SourceContent.

PDF path (PyMuPDF / fitz):
- Iterate text blocks per page, capture span font sizes.
- Body median = the most common rounded font size across the document.
- Headings = text spans whose size > body median * 1.18 AND length 2-200 chars.
- Images extracted via page.get_images() and reassembled with their pixmap data.

DOCX path (python-docx + raw ZIP):
- python-docx walks paragraphs, reading paragraph.style.name. "Heading 1/2/3" →
  heading, anything else → body text.
- Inline images live in ``word/media/*`` inside the .docx ZIP. We surface every
  image-bytes blob plus the first one as the brand-logo candidate (common
  enough that the logo is on the cover page / first inline image).

Both paths emit a ParsedDocument carrying:
  - raw_text (concatenated body)
  - headings (deduped, length-capped)
  - images (raw bytes, ordered as found)
  - title (document title metadata, falls back to first heading)
"""

from __future__ import annotations

import base64
import hashlib
import io
import logging
from collections import Counter
from dataclasses import dataclass, field
from typing import Literal

import docx as python_docx  # python-docx
import fitz  # PyMuPDF — `import fitz` is correct for the pymupdf package
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# Image extension → MIME, shared by the PDF and DOCX paths.
_IMG_EXT_TO_MIME = {
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "gif": "image/gif",
    "webp": "image/webp",
    "bmp": "image/bmp",
    "tiff": "image/tiff",
    "tif": "image/tiff",
}


SourceKindDoc = Literal["pdf", "docx"]


@dataclass
class OutlineBlock:
    """One block of the document in reading order.

    ``level`` is the heading level: ``1``/``2``/``3`` for headings (1 = most
    prominent), ``0`` for body text. Preserving order + level is what lets
    ``doc_structure`` group body copy under the right title — the flat
    ``raw_text``/``headings`` fields lose that association.
    """

    level: int  # 0 = body, 1/2/3 = heading levels
    text: str


@dataclass
class DocImage:
    """An image extracted in document order.

    ``anchor`` is ``len(outline)`` at the moment the image was encountered, i.e.
    it sits *after* ``outline[anchor - 1]``. That lets ``doc_structure`` tie each
    image to the page (and nearest heading) it appears under, instead of dumping
    every image on the homepage. ``width``/``height`` (pixels) drive both quality
    filtering and the matcher's size bonus.
    """

    data: bytes
    mime: str
    width: int | None = None
    height: int | None = None
    anchor: int = 0


@dataclass
class ParsedDocument:
    source_kind: SourceKindDoc
    source_ref: str           # original filename
    title: str | None
    raw_text: str
    headings: list[str] = field(default_factory=list)
    images: list[DocImage] = field(default_factory=list)  # in document order, w/ dims + anchor
    outline: list[OutlineBlock] = field(default_factory=list)  # ordered blocks w/ heading level


# --- PDF ------------------------------------------------------------------------


_HEADING_MIN_SCALE = 1.18  # font-size > body_median * this counts as a heading
_HEADING_MIN_LEN = 2
_HEADING_MAX_LEN = 200


def parse_pdf(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Extract text + headings + images from a PDF. Requires a text layer —
    image-only / scanned PDFs raise. Callers should check raw_text length and
    surface a helpful message."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"Could not open PDF: {exc}") from exc

    # Pass 1: collect font sizes to determine the body median
    sizes: list[int] = []
    for page in doc:
        for block in page.get_text("dict")["blocks"]:
            if block.get("type") != 0:  # 0 = text block
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    size = span.get("size")
                    if isinstance(size, (int, float)) and span.get("text", "").strip():
                        sizes.append(round(size))

    if not sizes:
        # PDF has no text layer (image-only / scanned). Caller decides how to surface.
        title = _pdf_title(doc)
        doc.close()
        return ParsedDocument(
            source_kind="pdf",
            source_ref=filename,
            title=title,
            raw_text="",
            headings=[],
            images=[],
        )

    body_median = Counter(sizes).most_common(1)[0][0]
    heading_threshold = body_median * _HEADING_MIN_SCALE

    # PDFs carry no heading *levels* — only font sizes. Rank the distinct heading
    # sizes (largest → H1, next → H2, rest → H3) so the outline can reconstruct a
    # hierarchy the way DOCX styles give us for free.
    heading_sizes = sorted(
        {round(s) for s in sizes if s >= heading_threshold}, reverse=True
    )

    # Pass 2: pull text + headings + ordered outline + interleaved images, in
    # document order. TEXT_PRESERVE_IMAGES surfaces image blocks (type 1) in the
    # same block stream as text — carrying their bytes + pixel dimensions — so
    # each image's position relative to the headings is preserved for per-page
    # association downstream.
    text_chunks: list[str] = []
    headings: list[str] = []
    outline: list[OutlineBlock] = []
    images: list[DocImage] = []
    seen_headings: set[str] = set()
    seen_image_hashes: set[bytes] = set()
    dict_flags = fitz.TEXTFLAGS_DICT | fitz.TEXT_PRESERVE_IMAGES
    for page in doc:
        for block in page.get_text("dict", flags=dict_flags)["blocks"]:
            btype = block.get("type")
            if btype == 1:  # image block
                _collect_pdf_image(block, images, seen_image_hashes, anchor=len(outline))
                continue
            if btype != 0:
                continue
            for line in block.get("lines", []):
                line_text_parts: list[str] = []
                line_max_size = 0.0
                for span in line.get("spans", []):
                    span_text = (span.get("text") or "").strip()
                    if not span_text:
                        continue
                    line_text_parts.append(span_text)
                    line_max_size = max(line_max_size, float(span.get("size") or 0))
                if not line_text_parts:
                    continue
                line_text = " ".join(line_text_parts)
                text_chunks.append(line_text)
                is_heading = (
                    line_max_size >= heading_threshold
                    and _HEADING_MIN_LEN <= len(line_text) <= _HEADING_MAX_LEN
                )
                if is_heading:
                    outline.append(
                        OutlineBlock(
                            level=_pdf_size_to_level(line_max_size, heading_sizes),
                            text=line_text,
                        )
                    )
                    if line_text.lower() not in seen_headings:
                        seen_headings.add(line_text.lower())
                        headings.append(line_text)
                else:
                    outline.append(OutlineBlock(level=0, text=line_text))

    raw_text = "\n".join(text_chunks)

    title = _pdf_title(doc) or (headings[0] if headings else None)
    doc.close()

    return ParsedDocument(
        source_kind="pdf",
        source_ref=filename,
        title=title,
        raw_text=raw_text,
        headings=headings[:50],
        images=images,
        outline=outline,
    )


def _collect_pdf_image(
    block: dict,
    images: list[DocImage],
    seen_hashes: set[bytes],
    *,
    anchor: int,
) -> None:
    """Append a dict image block as a :class:`DocImage`, de-duped by content.

    The same logo/photo repeated on every page shares identical bytes, so we
    hash to keep only the first occurrence (and its document position)."""
    data = block.get("image")
    if not data:
        return
    digest = hashlib.sha1(data).digest()
    if digest in seen_hashes:
        return
    seen_hashes.add(digest)
    ext = (block.get("ext") or "png").lower()
    images.append(
        DocImage(
            data=data,
            mime=_IMG_EXT_TO_MIME.get(ext, "image/png"),
            width=block.get("width"),
            height=block.get("height"),
            anchor=anchor,
        )
    )


def _pdf_size_to_level(size: float, heading_sizes: list[int]) -> int:
    """Map a heading's font size to a 1-3 level using the document's size tiers.

    ``heading_sizes`` is sorted largest-first. The largest distinct size is H1,
    the next is H2, everything smaller collapses to H3 so deeply varied PDFs
    don't explode into many levels. Falls back to level 1 if the size isn't a
    known tier (shouldn't happen, but keeps a heading from becoming body).
    """
    rounded = round(size)
    for idx, tier in enumerate(heading_sizes[:3]):
        if rounded >= tier:
            return idx + 1
    return min(len(heading_sizes), 3) or 1


def _pdf_title(doc) -> str | None:
    """Read /Title metadata; fall back to None."""
    try:
        meta = doc.metadata or {}
        title = (meta.get("title") or "").strip()
        return title or None
    except Exception:  # noqa: BLE001
        return None


# --- DOCX -----------------------------------------------------------------------


_HEADING_STYLE_PREFIXES = ("heading", "title", "subtitle")
_DOCX_BLIP = qn("a:blip")
_DOCX_EMBED = qn("r:embed")


def parse_docx(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Extract text + headings + images from a DOCX file."""
    try:
        document = python_docx.Document(io.BytesIO(file_bytes))
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"Could not open DOCX: {exc}") from exc

    text_parts: list[str] = []
    headings: list[str] = []
    outline: list[OutlineBlock] = []
    images: list[DocImage] = []
    seen_headings: set[str] = set()
    seen_image_hashes: set[bytes] = set()
    title_from_style: str | None = None

    # Walk paragraphs in document order. Inline images (<w:drawing><a:blip>) are
    # collected at their paragraph's position so each one stays associated with
    # the heading it sits under — an empty paragraph that holds only an image
    # still contributes the image at the right anchor.
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        blips = paragraph._p.findall(".//" + _DOCX_BLIP)
        if not text and not blips:
            continue
        if text:
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            level = _docx_style_to_level(style_name)
            if level > 0:
                outline.append(OutlineBlock(level=level, text=text))
                if _HEADING_MIN_LEN <= len(text) <= _HEADING_MAX_LEN:
                    key = text.lower()
                    if key not in seen_headings:
                        seen_headings.add(key)
                        headings.append(text)
                if title_from_style is None and style_name.startswith("title"):
                    title_from_style = text
            else:
                outline.append(OutlineBlock(level=0, text=text))
            text_parts.append(text)
        for blip in blips:
            _collect_docx_image(
                blip, document, images, seen_image_hashes, anchor=len(outline)
            )

    # Tables — flatten cell text so the LLM sees pricing tables, hours, etc.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                row_text = " | ".join(cells)
                text_parts.append(row_text)
                outline.append(OutlineBlock(level=0, text=row_text))

    raw_text = "\n".join(text_parts)

    title = title_from_style or _docx_core_title(document) or (headings[0] if headings else None)

    return ParsedDocument(
        source_kind="docx",
        source_ref=filename,
        title=title,
        raw_text=raw_text,
        headings=headings[:50],
        images=images,
        outline=outline,
    )


def _collect_docx_image(
    blip,
    document,
    images: list[DocImage],
    seen_hashes: set[bytes],
    *,
    anchor: int,
) -> None:
    """Resolve a ``<a:blip>`` to its image part and append it as a DocImage."""
    rid = blip.get(_DOCX_EMBED)
    if not rid:
        return
    part = document.part.related_parts.get(rid)
    if part is None:
        return
    try:
        data = part.blob
    except Exception as exc:  # noqa: BLE001
        logger.debug("Failed reading DOCX image %s: %s", rid, exc)
        return
    if not data:
        return
    digest = hashlib.sha1(data).digest()
    if digest in seen_hashes:
        return
    seen_hashes.add(digest)
    img = getattr(part, "image", None)
    images.append(
        DocImage(
            data=data,
            mime=getattr(part, "content_type", None) or "image/png",
            width=getattr(img, "px_width", None),
            height=getattr(img, "px_height", None),
            anchor=anchor,
        )
    )


def _docx_style_to_level(style_name: str) -> int:
    """Map a (lowercased) DOCX paragraph style to an outline level.

    ``Title`` → 1 and ``Subtitle`` → 2 so a cover-page title/tagline anchors the
    hierarchy; ``Heading N`` maps to ``min(N, 3)``; everything else is body (0).
    """
    if style_name.startswith("title"):
        return 1
    if style_name.startswith("subtitle"):
        return 2
    if style_name.startswith("heading"):
        digits = "".join(ch for ch in style_name if ch.isdigit())
        if digits:
            return min(int(digits), 3)
        return 1
    return 0


def _docx_core_title(document) -> str | None:
    try:
        props = document.core_properties
        title = (getattr(props, "title", "") or "").strip()
        return title or None
    except Exception:  # noqa: BLE001
        return None


# --- dispatcher -----------------------------------------------------------------


class DocParseError(Exception):
    """Raised when a document is unparseable. Carries a friendly status code."""

    def __init__(self, message: str, status: int = 400):
        super().__init__(message)
        self.status = status


def parse_document(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Dispatch by extension. Caller is responsible for catching DocParseError."""
    if not filename or not file_bytes:
        raise DocParseError("Empty file or missing filename", status=400)

    lower = filename.lower()
    try:
        if lower.endswith(".pdf"):
            parsed = parse_pdf(file_bytes, filename)
        elif lower.endswith(".docx"):
            parsed = parse_docx(file_bytes, filename)
        elif lower.endswith(".doc"):
            raise DocParseError(
                "Legacy .doc files aren't supported — please save as .docx.",
                status=415,
            )
        else:
            raise DocParseError(
                f"Unsupported file type: {filename}. Upload a PDF or DOCX.",
                status=415,
            )
    except ValueError as exc:
        raise DocParseError(str(exc), status=400) from exc

    if not parsed.raw_text.strip() or len(parsed.raw_text.strip()) < 80:
        raise DocParseError(
            "Could not extract enough text content from that file. "
            "If it's a scanned PDF, the generator needs a text layer — "
            "OCR isn't supported. Paste the content into the URL tab's textarea instead.",
            status=422,
        )

    return parsed


# --- helpers for the router -----------------------------------------------------


def image_to_data_url(image_bytes: bytes, mime: str) -> str:
    encoded = base64.b64encode(image_bytes).decode("ascii")
    return f"data:{mime};base64,{encoded}"
