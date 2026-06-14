"""
Render a planned site into an editable ``.docx`` content brief.

This is the inverse of ``doc_structure``: each page becomes a **Heading 1**, each
in-page heading a **Heading 2**, and the copy plain paragraphs — exactly the
shape ``parse_docx`` + ``split_into_pages`` read back. So a brief exported from a
scrape can be edited and re-uploaded to regenerate the site *from the document*.

v1 scope (see DOC_DRIVEN_GENERATION_PLAN.md): ``.docx`` only; images are
referenced as captioned URL placeholders rather than embedded.
"""

from __future__ import annotations

import io

import docx as python_docx

from app.models.content_blocks import SourceContent
from app.models.industry import PageScaffold
from app.services.source_router import match_scaffolds_to_pages

_MAX_IMAGES_PER_PAGE = 6
_INTRO = (
    "Each Heading 1 below is a page. Edit the titles and copy, then re-upload "
    "this document to regenerate the site from it. Sections and layout are "
    "rebuilt by the generator — you don't need to format them here."
)


def build_site_document(
    source: SourceContent,
    scaffolds: list[PageScaffold],
    *,
    site_name: str | None = None,
) -> bytes:
    """Return ``.docx`` bytes for the given source + inferred page scaffolds."""
    doc = python_docx.Document()
    doc.add_heading(site_name or source.title or "Website content", level=0)
    doc.add_paragraph(_INTRO)

    matched = match_scaffolds_to_pages(scaffolds, source)
    for scaffold in scaffolds:
        if scaffold.is_legal:
            # Legal pages are generated from boilerplate, not document copy.
            continue
        doc.add_heading(scaffold.title, level=1)
        page = matched.get(scaffold.slug)
        if page is not None:
            _render_page_body(doc, page, page_title=scaffold.title)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _render_page_body(doc, page: SourceContent, *, page_title: str) -> None:
    """Emit a page's copy, promoting its own headings to Heading 2.

    The extractor keeps each heading on its own line inside ``raw_text``, so we
    re-detect them via ``page.headings`` and lift them to Heading 2 — giving the
    brief a real title hierarchy that survives the round-trip.
    """
    heading_keys = {h.strip().lower() for h in (page.headings or []) if h.strip()}
    title_key = page_title.strip().lower()

    for line in (page.raw_text or "").split("\n"):
        text = line.strip()
        if not text:
            continue
        if text.lower() == title_key:
            continue  # already emitted as the page's Heading 1
        if text.lower() in heading_keys:
            doc.add_heading(text, level=2)
        else:
            doc.add_paragraph(text)

    for url in (page.images or [])[:_MAX_IMAGES_PER_PAGE]:
        caption = doc.add_paragraph(f"[image] {url}")
        for run in caption.runs:
            run.italic = True
